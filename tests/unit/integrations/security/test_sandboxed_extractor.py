"""Unit tests for Ephemeral Sandboxed Attachment Extractor (ADR-002 / ADR-003)."""

from __future__ import annotations

import io
import zipfile
from collections.abc import AsyncIterator

import pytest
from docx import Document

from cowork_agent.features.email_action_plan.schemas import ExtractionLimits
from cowork_agent.integrations.security.fakes import FakeAttachmentExtractor
from cowork_agent.integrations.security.sandboxed_extractor import EphemeralSandboxedExtractor


async def _to_async_stream(data: bytes, chunk_size: int = 1024) -> AsyncIterator[bytes]:
    """Helper to convert raw bytes into an async chunk iterator."""
    for i in range(0, len(data), chunk_size):
        yield data[i : i + chunk_size]


@pytest.mark.asyncio
async def test_sandboxed_extractor_text_document():
    extractor = EphemeralSandboxedExtractor()
    content = b"Line 1: Project milestone update.\nLine 2: Budget approved."
    stream = _to_async_stream(content)
    limits = ExtractionLimits(max_bytes=1024 * 1024, max_characters=1000, timeout_seconds=5)

    result = await extractor.extract(
        attachment_id="att-txt-01",
        filename="notes.txt",
        declared_mime_type="text/plain",
        content=stream,
        limits=limits,
    )

    assert result.status == "ok"
    assert result.warning_code is None
    assert result.text is not None
    assert "Project milestone update" in result.text
    assert len(result.units) == 1
    assert result.units[0].kind == "text_document"
    assert result.attachment_id == "att-txt-01"


@pytest.mark.asyncio
async def test_sandboxed_extractor_docx_document():
    # Generate in-memory DOCX
    doc = Document()
    doc.add_heading("Bao cao tien do", level=1)
    doc.add_paragraph("Doan van ban mo dau bao cao thang 8.")
    docx_buf = io.BytesIO()
    doc.save(docx_buf)
    docx_bytes = docx_buf.getvalue()

    extractor = EphemeralSandboxedExtractor()
    stream = _to_async_stream(docx_bytes)
    limits = ExtractionLimits(max_bytes=5 * 1024 * 1024, max_characters=5000, timeout_seconds=5)

    result = await extractor.extract(
        attachment_id="att-docx-01",
        filename="Bao_cao.docx",
        declared_mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        content=stream,
        limits=limits,
    )

    assert result.status == "ok"
    assert result.warning_code is None
    assert result.text is not None
    assert "Doan van ban mo dau" in result.text
    assert len(result.units) >= 1


@pytest.mark.asyncio
async def test_sandboxed_extractor_hard_size_limit():
    extractor = EphemeralSandboxedExtractor(hard_max_bytes=1024)  # 1 KB hard limit
    large_content = b"A" * 5000  # 5 KB
    stream = _to_async_stream(large_content, chunk_size=512)
    limits = ExtractionLimits(max_bytes=1024, max_characters=1000, timeout_seconds=5)

    result = await extractor.extract(
        attachment_id="att-large-01",
        filename="huge_file.txt",
        declared_mime_type="text/plain",
        content=stream,
        limits=limits,
    )

    assert result.status == "size_exceeded"
    assert result.warning_code == "SIZE_EXCEEDED"
    assert result.text is None


@pytest.mark.asyncio
async def test_sandboxed_extractor_quarantines_prohibited_script():
    extractor = EphemeralSandboxedExtractor()
    malicious_script = b"WScript.Echo \"Malicious VBS execution\""
    stream = _to_async_stream(malicious_script)
    limits = ExtractionLimits(max_bytes=1024 * 1024, max_characters=1000, timeout_seconds=5)

    result = await extractor.extract(
        attachment_id="att-vbs-01",
        filename="invoice.pdf.vbs",
        declared_mime_type="application/octet-stream",
        content=stream,
        limits=limits,
    )

    assert result.status == "quarantined"
    assert result.warning_code == "QUARANTINED_MALWARE"
    assert result.text is None


@pytest.mark.asyncio
async def test_sandboxed_extractor_quarantines_zip_bomb():
    extractor = EphemeralSandboxedExtractor()
    # Create zip bomb in memory (high compression ratio)
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("zeroes.bin", b"0" * (5 * 1024 * 1024))  # 5MB of zeroes compresses to < 5KB
    zip_bytes = buf.getvalue()

    stream = _to_async_stream(zip_bytes)
    limits = ExtractionLimits(max_bytes=10 * 1024 * 1024, max_characters=1000, timeout_seconds=5)

    result = await extractor.extract(
        attachment_id="att-bomb-01",
        filename="archive.zip",
        declared_mime_type="application/zip",
        content=stream,
        limits=limits,
    )

    assert result.status == "quarantined"
    assert result.warning_code == "QUARANTINED_ZIP_BOMB"
    assert result.text is None


@pytest.mark.asyncio
async def test_sandboxed_extractor_unsupported_format():
    extractor = EphemeralSandboxedExtractor()
    unknown_binary = b"\x00\x01\x02\x03RANDOM_PROPRIETARY_DATA"
    stream = _to_async_stream(unknown_binary)
    limits = ExtractionLimits(max_bytes=1024 * 1024, max_characters=1000, timeout_seconds=5)

    result = await extractor.extract(
        attachment_id="att-raw-01",
        filename="data.proprietary_raw_xyz",
        declared_mime_type="application/octet-stream",
        content=stream,
        limits=limits,
    )

    assert result.status == "unsupported"
    assert result.warning_code == "UNSUPPORTED_FORMAT"
    assert result.text is None


@pytest.mark.asyncio
async def test_fake_attachment_extractor():
    fake = FakeAttachmentExtractor(fixed_text="Pre-configured test text")
    stream = _to_async_stream(b"input data")
    limits = ExtractionLimits(max_bytes=1024 * 1024, max_characters=1000, timeout_seconds=5)

    result = await fake.extract(
        attachment_id="fake-01",
        filename="test.txt",
        declared_mime_type="text/plain",
        content=stream,
        limits=limits,
    )

    assert result.status == "ok"
    assert result.text == "Pre-configured test text"
    assert len(fake.extracted_attachments) == 1
