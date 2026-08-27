from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path
from xml.etree import ElementTree as ET
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document

from cowork_agent.config import KnowledgeIngestionSettings
from cowork_agent.integrations.knowledge_ingestion.docx_extractor import DocxExtractor
from cowork_agent.integrations.knowledge_ingestion.models import PdfInspection, PdfKind
from cowork_agent.integrations.knowledge_ingestion.service import KnowledgeIngestionService
from cowork_agent.integrations.knowledge_ingestion.text_sanitizer import (
    split_frontmatter,
)

_NFD_VIETNAMESE_E = "e\u0302\u0301"
_MANIFEST_NAME = "ingestion-manifest.json"
_DCTERMS_CREATED = "{http://purl.org/dc/terms/}created"
_DCTERMS_MODIFIED = "{http://purl.org/dc/terms/}modified"


class StubPdfInspector:
    def __init__(self, inspection: PdfInspection) -> None:
        self.inspection = inspection
        self.paths: list[Path] = []

    def inspect(self, path: Path) -> PdfInspection:
        self.paths.append(path)
        return self.inspection


def _settings() -> KnowledgeIngestionSettings:
    return KnowledgeIngestionSettings.from_env({"KNOWLEDGE_INGEST_OCR_ENABLED": "false"})


def _write_docx(path: Path, text: str = "Knowledge body") -> None:
    document = Document()
    document.add_heading("Policy", level=1)
    document.add_paragraph(text)
    document.save(path)


def _drop_docx_core_dates(path: Path) -> None:
    with ZipFile(path) as archive:
        contents = {info.filename: archive.read(info.filename) for info in archive.infolist()}
    core = ET.fromstring(contents["docProps/core.xml"])
    drop_tags = {_DCTERMS_CREATED, _DCTERMS_MODIFIED}
    for child in list(core):
        if child.tag in drop_tags:
            core.remove(child)
    contents["docProps/core.xml"] = ET.tostring(core, encoding="utf-8", xml_declaration=True)
    with ZipFile(path, "w", compression=ZIP_DEFLATED) as archive:
        for name, data in contents.items():
            archive.writestr(name, data)


def _native_pdf() -> PdfInspection:
    return PdfInspection(PdfKind.TEXT_BASED, 1, (), {1: "native"})


class _StubOcrExtractor:
    def __init__(self, output: str = "# OCR Result") -> None:
        self.output = output
        self.extracted_files: list[str] = []

    def extract(self, filename: str, content: bytes) -> str:
        self.extracted_files.append(filename)
        return self.output


def test_service_docx_and_pdf_native_ingestion_and_manifest_reuse(tmp_path: Path) -> None:
    raw = tmp_path / "raw"
    raw.mkdir()
    _write_docx(raw / "policy.docx")
    (raw / "Policy File.pdf").write_bytes(b"pdf")
    (raw / "note.txt").write_text(
        f"# Note\n\nHello {_NFD_VIETNAMESE_E}\x00world\n", encoding="utf-8"
    )
    (raw / "markdown_doc.md").write_text(
        "---\ntitle: Old\n---\n# Real Title\nBody\n", encoding="utf-8"
    )

    pdf_inspection = PdfInspection(PdfKind.TEXT_BASED, 2, (), {1: "One", 2: "Two"})
    service = KnowledgeIngestionService(
        _settings(), DocxExtractor(), StubPdfInspector(pdf_inspection)
    )

    # First ingest -> succeeded
    first_outcomes = service.ingest(raw, tmp_path / "extracted", force=False)
    assert all(o.status == "succeeded" for o in first_outcomes)

    # Re-ingest -> skipped
    second_outcomes = service.ingest(raw, tmp_path / "extracted", force=False)
    assert all(o.status == "skipped" for o in second_outcomes)

    # Verify DOCX frontmatter & body
    docx_md = (tmp_path / "extracted" / "policy.md").read_text(encoding="utf-8")
    d_fields, d_body = split_frontmatter(docx_md)
    assert d_fields["title"] == "Policy" and "# Policy" in d_body

    # Verify PDF page markers
    pdf_md = (tmp_path / "extracted" / "policy-file.md").read_text(encoding="utf-8")
    assert "<!-- Page 1 -->" in pdf_md and "<!-- Page 2 -->" in pdf_md

    # Verify TXT sanitization (NFC normalized, null byte stripped)
    txt_md = (tmp_path / "extracted" / "note.md").read_text(encoding="utf-8")
    t_fields, t_body = split_frontmatter(txt_md)
    assert t_fields["extractor"] == "text" and "\u1ebf" in t_body and "\x00" not in t_body

    # Verify Markdown frontmatter not nested
    md_md = (tmp_path / "extracted" / "markdown-doc.md").read_text(encoding="utf-8")
    m_fields, m_body = split_frontmatter(md_md)
    assert m_fields["extractor"] == "markdown" and not m_body.lstrip().startswith("---")


def test_service_ocr_extraction_modes_adaptive_and_advance(tmp_path: Path) -> None:
    raw = tmp_path / "raw"
    raw.mkdir()
    (raw / "clean.pdf").write_bytes(b"clean-pdf")
    (raw / "mixed.pdf").write_bytes(b"mixed-pdf")
    (raw / "advance.pdf").write_bytes(b"advance-pdf")

    settings = KnowledgeIngestionSettings.from_env(
        {"EXTRACTION_MODE": "adaptive", "MISTRAL_API_KEY": "secret"}
    )
    stub_ocr = _StubOcrExtractor("# OCR Content")

    def inspect_fn(path: Path) -> PdfInspection:
        if path.name == "clean.pdf":
            return PdfInspection(PdfKind.TEXT_BASED, 1, (), {1: "clean digital text"})
        return PdfInspection(PdfKind.MIXED, 2, (2,), {1: "page 1"})

    inspector = StubPdfInspector(PdfInspection(PdfKind.TEXT_BASED, 1, (), {}))
    inspector.inspect = inspect_fn  # type: ignore[assignment]

    service = KnowledgeIngestionService(
        settings, DocxExtractor(), inspector, ocr_extractor=stub_ocr
    )
    outcomes = service.ingest(raw, tmp_path / "extracted", force=False)
    assert all(o.status == "succeeded" for o in outcomes)
    assert "mixed.pdf" in stub_ocr.extracted_files
    assert "clean.pdf" not in stub_ocr.extracted_files


def test_service_validation_errors_and_limit_guards(tmp_path: Path) -> None:
    raw = tmp_path / "raw"
    raw.mkdir()
    _write_docx(raw / "A B.docx")
    _write_docx(raw / "A-B.docx")
    (raw / "bad.txt").write_bytes(b"\xff\xfe not utf-8")
    (raw / "long.pdf").write_bytes(b"pdf")

    settings = KnowledgeIngestionSettings.from_env(
        {
            "KNOWLEDGE_INGEST_OCR_ENABLED": "false",
            "KNOWLEDGE_INGEST_MAX_PDF_PAGES": "1",
        }
    )
    service = KnowledgeIngestionService(
        settings,
        DocxExtractor(),
        StubPdfInspector(PdfInspection(PdfKind.TEXT_BASED, 2, (), {1: "a", 2: "b"})),
    )

    outcomes = service.ingest(raw, tmp_path / "extracted", force=False)
    reasons = {o.source: o.reason_code for o in outcomes}
    assert reasons["A B.docx"] == "output_name_collision"
    assert reasons["bad.txt"] == "decode_failed"
    assert reasons["long.pdf"] == "pdf_page_limit_exceeded"


def test_service_date_harvesting_and_manifest_sync(tmp_path: Path) -> None:
    raw = tmp_path / "raw"
    raw.mkdir()
    document = Document()
    document.add_heading("Dated Policy", level=1)
    document.add_paragraph("Knowledge body")
    document.core_properties.created = datetime(2026, 8, 7, 10, 19, 25)
    document.save(raw / "dated.docx")

    _write_docx(raw / "undated.docx")
    _drop_docx_core_dates(raw / "undated.docx")

    extracted = tmp_path / "extracted"
    service = KnowledgeIngestionService(
        _settings(), DocxExtractor(), StubPdfInspector(_native_pdf())
    )
    outcomes = service.ingest(raw, extracted, force=False)
    assert all(o.status == "succeeded" for o in outcomes)

    manifest = json.loads((extracted / _MANIFEST_NAME).read_text(encoding="utf-8"))
    assert manifest["dated.docx"]["document_date"] == "2026-08-07"
    assert manifest["undated.docx"]["document_date"] == ""
