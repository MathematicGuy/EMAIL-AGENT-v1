from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from cowork_agent.integrations.knowledge_ingestion.docx_extractor import DocxExtractor

pytestmark = pytest.mark.extended


def test_docx_extractor_preserves_headings_lists_and_tables_in_document_order(
    tmp_path: Path,
) -> None:
    """Removing structural conversions would lose navigable document meaning."""
    document = Document()
    document.add_heading("Quy định", level=1)
    document.add_paragraph("Nội dung mở đầu")
    document.add_heading("Hồ sơ", level=2)
    document.add_paragraph("Giấy tờ", style="List Bullet")
    document.add_heading("Chi tiết", level=3)
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Giấy tờ"
    table.rows[0].cells[1].text = "Bắt buộc"
    table.rows[1].cells[0].text = "A|B"
    table.rows[1].cells[1].text = "Có"
    path = tmp_path / "policy.docx"
    document.save(path)

    result = DocxExtractor().extract(path)

    assert result.markdown == (
        "# Quy định\n\n"
        "Nội dung mở đầu\n\n"
        "## Hồ sơ\n\n"
        "- Giấy tờ\n\n"
        "### Chi tiết\n\n"
        "| Giấy tờ | Bắt buộc |\n"
        "| --- | --- |\n"
        "| A\\|B | Có |"
    )
    assert result.page_count == 1


def _bold_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.add_run(text).bold = True


def test_fully_bold_paragraph_becomes_a_heading_at_its_structural_depth(
    tmp_path: Path,
) -> None:
    """Vietnamese statutes style every paragraph Normal and mark structure with
    bold alone; extraction is the last stage that can see it."""
    document = Document()
    _bold_paragraph(document, "Chương I")
    _bold_paragraph(document, "Điều 1. Phạm vi điều chỉnh")
    document.add_paragraph("Luật này quy định về chế độ sở hữu đất đai.")
    path = tmp_path / "statute.docx"
    document.save(path)

    result = DocxExtractor().extract(path)

    assert result.markdown == (
        "## Chương I\n\n"
        "#### Điều 1. Phạm vi điều chỉnh\n\n"
        "Luật này quy định về chế độ sở hữu đất đai."
    )


def test_bold_paragraph_without_structural_keyword_uses_the_default_depth(
    tmp_path: Path,
) -> None:
    document = Document()
    _bold_paragraph(document, "QUY ĐỊNH CHUNG")
    path = tmp_path / "titled.docx"
    document.save(path)

    assert DocxExtractor().extract(path).markdown == "### QUY ĐỊNH CHUNG"


def test_sentence_containing_a_bold_phrase_stays_a_paragraph(tmp_path: Path) -> None:
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("Điều 1. ").bold = True
    paragraph.add_run("Phạm vi điều chỉnh của luật này rất rộng.")
    path = tmp_path / "mixed.docx"
    document.save(path)

    assert DocxExtractor().extract(path).markdown == (
        "Điều 1. Phạm vi điều chỉnh của luật này rất rộng."
    )


def test_long_bold_run_on_paragraph_stays_a_paragraph(tmp_path: Path) -> None:
    prose = "Bước 1: Cá nhân chuẩn bị hồ sơ theo quy định của pháp luật. " * 8
    document = Document()
    _bold_paragraph(document, prose.strip())
    path = tmp_path / "runon.docx"
    document.save(path)

    assert DocxExtractor().extract(path).markdown == prose.strip()


def test_word_heading_style_still_wins_over_bold_inference(tmp_path: Path) -> None:
    document = Document()
    heading = document.add_heading("Điều 1. Phạm vi", level=1)
    for run in heading.runs:
        run.bold = True
    path = tmp_path / "styled.docx"
    document.save(path)

    assert DocxExtractor().extract(path).markdown == "# Điều 1. Phạm vi"
