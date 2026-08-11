from __future__ import annotations

from pathlib import Path

from docx import Document

from cowork_agent.config import KnowledgeIngestionSettings
from cowork_agent.integrations.knowledge_ingestion.docx_extractor import DocxExtractor
from cowork_agent.integrations.knowledge_ingestion.models import PdfInspection, PdfKind
from cowork_agent.integrations.knowledge_ingestion.service import KnowledgeIngestionService
from cowork_agent.integrations.rag.knowledge_base import load_corpus


class NativePdfInspector:
    def inspect(self, path: Path) -> PdfInspection:
        return PdfInspection(PdfKind.TEXT_BASED, 1, (), {1: "unused"})


def test_ingested_docx_markdown_is_loadable_by_rag(tmp_path: Path) -> None:
    """Breaking emitted Markdown would make the administrator corpus unusable by RAG."""
    source = tmp_path / "raw"
    source.mkdir()
    document = Document()
    document.add_heading("Expense policy", level=1)
    document.add_paragraph("Submit receipts within five days.")
    document.save(source / "expense-policy.docx")
    settings = KnowledgeIngestionSettings.from_env(
        {"KNOWLEDGE_INGEST_OCR_ENABLED": "false"}, load_env_file=False
    )
    service = KnowledgeIngestionService(settings, DocxExtractor(), NativePdfInspector())

    outcomes = service.ingest(source, tmp_path / "extracted", force=False)
    documents = load_corpus(tmp_path / "extracted", tenant_id="local")

    assert outcomes[0].status == "succeeded"
    assert documents[0].title == "Expense policy"
    assert documents[0].chunks[0].text == "Submit receipts within five days."
