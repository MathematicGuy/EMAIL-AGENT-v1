"""FastAPI REST router for Artifact Creation, Preview and Download Flow (ADR-008)."""

import asyncio
import io
import os
import re
from datetime import UTC, datetime
from pathlib import Path
from typing import Any
from urllib.parse import quote

import docx
from fastapi import APIRouter, HTTPException, Request, Response
from pydantic import BaseModel, Field

from cowork_agent.features.ai_chat.ports import (
    extract_stem_from_content,
    sanitize_filename,
)

router = APIRouter(prefix="/api/v1/reports", tags=["reports"])


class ReportCreate(BaseModel):
    filename: str = Field(..., description="Original requested filename")
    content: str = Field(..., description="Markdown content of the report")


class ReportFile(BaseModel):
    ref_id: str
    filename: str
    object_key: str
    created_at: str
    size_bytes: int


def _add_formatted_text_to_paragraph(paragraph: Any, text: str) -> None:
    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif (
            (part.startswith("*") and part.endswith("*") and len(part) >= 2)
            or (part.startswith("_") and part.endswith("_") and len(part) >= 2)
        ):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`") and len(part) >= 2:
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
        else:
            paragraph.add_run(part)


def markdown_to_docx_bytes(markdown_text: str) -> bytes:
    doc = docx.Document()
    for line in markdown_text.splitlines():
        line_str = line.strip()
        if not line_str:
            continue
        if line_str.startswith("# "):
            p = doc.add_heading(level=1)
            _add_formatted_text_to_paragraph(p, line_str[2:].strip())
        elif line_str.startswith("## "):
            p = doc.add_heading(level=2)
            _add_formatted_text_to_paragraph(p, line_str[3:].strip())
        elif line_str.startswith("### "):
            p = doc.add_heading(level=3)
            _add_formatted_text_to_paragraph(p, line_str[4:].strip())
        elif line_str.startswith("- ") or line_str.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            _add_formatted_text_to_paragraph(p, line_str[2:].strip())
        else:
            p = doc.add_paragraph()
            _add_formatted_text_to_paragraph(p, line_str)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


class LocalReportStorage:
    """Local filesystem storage for workspace report artifacts."""

    def __init__(self, workspace_dir: str | Path = "workspace/reports") -> None:
        self._workspace_dir = Path(workspace_dir)

    def _ensure_dir(self) -> Path:
        self._workspace_dir.mkdir(parents=True, exist_ok=True)
        return self._workspace_dir

    async def upload_bytes(
        self, object_key: str, data: bytes, content_type: str = "text/markdown"
    ) -> None:
        del content_type
        base_dir = self._ensure_dir()
        filename = Path(object_key).name
        file_path = base_dir / filename
        await asyncio.to_thread(file_path.write_bytes, data)

    def _find_file(self, object_key: str) -> Path | None:
        base_dir = self._ensure_dir()
        target = Path(object_key).name
        exact_path = base_dir / target
        if exact_path.exists() and exact_path.is_file():
            return exact_path

        target_lower = target.lower()
        for item in base_dir.iterdir():
            if item.is_file() and item.name.lower() == target_lower:
                return item

        target_stem = target.removesuffix(".md").lower()
        for item in base_dir.iterdir():
            if item.is_file() and item.name.removesuffix(".md").lower() == target_stem:
                return item
        return None

    async def download_bytes(self, object_key: str) -> bytes:
        file_path = self._find_file(object_key)
        if file_path is None:
            raise KeyError(f"Report file {object_key} not found")
        return await asyncio.to_thread(file_path.read_bytes)

    async def list_objects(self, prefix: str = "") -> list[dict[str, object]]:
        del prefix
        base_dir = self._ensure_dir()

        def _scan() -> list[dict[str, object]]:
            results: list[dict[str, object]] = []
            if not base_dir.exists():
                return results
            for item in base_dir.iterdir():
                if item.is_file():
                    stat = item.stat()
                    results.append(
                        {
                            "name": item.name,
                            "metadata": {"size": stat.st_size},
                            "created_at": datetime.fromtimestamp(
                                stat.st_mtime, UTC
                            ).isoformat(),
                        }
                    )
            return results

        return await asyncio.to_thread(_scan)

    async def delete(self, object_key: str) -> None:
        file_path = self._find_file(object_key)
        if file_path is not None:
            await asyncio.to_thread(file_path.unlink)


def _get_storage(request: Request) -> Any:
    storage = getattr(request.app.state, "report_storage", None)
    if storage is None:
        storage = LocalReportStorage()
        request.app.state.report_storage = storage
    return storage


@router.get("", response_model=list[ReportFile], operation_id="reports_list_reports")
async def list_reports(request: Request) -> list[ReportFile]:
    storage = _get_storage(request)

    try:
        objects = await storage.list_objects()
    except Exception as exc:
        raise HTTPException(status_code=503, detail="storage unavailable") from exc

    reports: list[ReportFile] = []
    for obj in objects:
        name = str(obj.get("name", ""))
        if not name or name.endswith("/"):
            continue
        ref_id = name
        filename = name

        metadata = obj.get("metadata") or {}
        size_bytes = int(metadata.get("size", 0)) if isinstance(metadata, dict) else 0
        created_at = str(obj.get("created_at", ""))

        reports.append(
            ReportFile(
                ref_id=ref_id,
                filename=filename,
                object_key=ref_id,
                created_at=created_at,
                size_bytes=size_bytes,
            )
        )
    reports.sort(key=lambda r: r.created_at or "", reverse=True)
    return reports


@router.post("", response_model=ReportFile, status_code=201, operation_id="reports_create_report")
async def create_report(request: Request, body: ReportCreate) -> ReportFile:
    storage = _get_storage(request)

    base_stem = extract_stem_from_content(body.content, body.filename)
    ref_id = f"{base_stem}.md"

    object_key = ref_id
    content_bytes = body.content.encode("utf-8")

    try:
        await storage.upload_bytes(object_key, content_bytes, content_type="text/markdown")
    except Exception as exc:
        raise HTTPException(status_code=503, detail="storage unavailable") from exc

    return ReportFile(
        ref_id=ref_id,
        filename=ref_id,
        object_key=object_key,
        created_at=datetime.now(UTC).isoformat(),
        size_bytes=len(content_bytes),
    )


@router.get("/{ref_id}", operation_id="reports_get_report_content")
async def get_report_content(request: Request, ref_id: str) -> Response:
    storage = _get_storage(request)

    safe_ref = os.path.basename(ref_id)

    try:
        data = await storage.download_bytes(safe_ref)
    except KeyError as exc:
        raise HTTPException(status_code=404, detail="File not found") from exc
    except Exception as exc:
        raise HTTPException(status_code=503, detail="storage unavailable") from exc

    return Response(content=data, media_type="text/markdown; charset=utf-8")


@router.delete("/{ref_id}", operation_id="reports_delete_report")
async def delete_report(request: Request, ref_id: str) -> dict[str, bool]:
    storage = _get_storage(request)

    safe_ref = os.path.basename(ref_id)

    try:
        await storage.delete(safe_ref)
    except Exception as exc:
        raise HTTPException(status_code=503, detail="storage unavailable") from exc

    return {"success": True}


@router.get("/{ref_id}/download", operation_id="reports_download_report_docx")
async def download_report_docx(request: Request, ref_id: str) -> Response:
    storage = _get_storage(request)

    safe_ref = os.path.basename(ref_id)

    try:
        md_bytes = await storage.download_bytes(safe_ref)
    except KeyError as exc:
        raise HTTPException(status_code=404, detail="File not found") from exc
    except Exception as exc:
        raise HTTPException(status_code=503, detail="storage unavailable") from exc

    md_text = md_bytes.decode("utf-8", errors="replace")
    docx_bytes = markdown_to_docx_bytes(md_text)

    docx_filename = os.path.splitext(safe_ref)[0] + ".docx"
    ascii_filename = re.sub(r"[^\x00-\x7F]+", "_", docx_filename)
    encoded_filename = quote(docx_filename)
    content_disposition = (
        f'attachment; filename="{ascii_filename}"; filename*=UTF-8\'\'{encoded_filename}'
    )

    return Response(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": content_disposition},
    )
